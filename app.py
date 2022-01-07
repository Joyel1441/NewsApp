from ast import parse
from flask import Flask, render_template, request, redirect, url_for
from flask.helpers import send_from_directory
import requests
from docx import Document
from openpyxl import Workbook
import requests
from bs4 import BeautifulSoup
import csv
import json
from flask_restful import Resource, Api, reqparse


app = Flask(__name__)
api = Api(app)

# index page
@app.route('/')
def Index():

    # define lists
    title = []
    link = []
    description = []
    date = []
    img_url = []
    source = []
   
    # get response
    response = requests.get("https://newsdata.io/api/1/news?apikey=pub_3357cb92f61b1a3819ff35e2ee398a56d2b3&language=en")
    json_response = response.json()

    # get news data
    for i in range(len(json_response['results'])) :
        try:
            if json_response['results'][i]['full_description'] != None :
                description.append(json_response['results'][i]['full_description'])
            elif json_response['results'][i]['description'] != None :
                description.append(json_response['results'][i]['description'])
            else:
                continue
        except:
            continue
        title.append(json_response['results'][i]['title'])
        link.append(json_response['results'][i]['link'])
        date.append(json_response['results'][i]['pubDate'])
        img_url.append(json_response['results'][i]['image_url'])
        source.append(json_response['results'][i]['source_id'])

    data = zip(title, link, description, date, img_url, source)

    return render_template('index.html', data=data)

# technology page
@app.route('/technology')
def technology():
    title = []
    link = []
    description = []
    date = []
    img_url = []
    source = []

    response = requests.get("https://newsdata.io/api/1/news?apikey=pub_3357cb92f61b1a3819ff35e2ee398a56d2b3&language=en&category=technology")
    json_response = response.json()

    for i in range(len(json_response['results'])) :
        try:
            if json_response['results'][i]['full_description'] != None :
                description.append(json_response['results'][i]['full_description'])
            elif json_response['results'][i]['description'] != None :
                description.append(json_response['results'][i]['description'])
            else:
                continue
        except:
            continue
        title.append(json_response['results'][i]['title'])
        link.append(json_response['results'][i]['link'])
        date.append(json_response['results'][i]['pubDate'])
        img_url.append(json_response['results'][i]['image_url'])
        source.append(json_response['results'][i]['source_id'])

    data = zip(title, link, description, date, img_url, source)

    return render_template('technology.html', data=data)

# sports page
@app.route('/sports')
def sports():
    title = []
    link = []
    description = []
    date = []
    img_url = []
    source = []

    response = requests.get("https://newsdata.io/api/1/news?apikey=pub_3357cb92f61b1a3819ff35e2ee398a56d2b3&language=en&category=sports")
    json_response = response.json()

    for i in range(len(json_response['results'])) :
        try:
            if json_response['results'][i]['full_description'] != None :
                description.append(json_response['results'][i]['full_description'])
            elif json_response['results'][i]['description'] != None :
                description.append(json_response['results'][i]['description'])
            else:
                continue
        except:
            continue
        title.append(json_response['results'][i]['title'])
        link.append(json_response['results'][i]['link'])
        date.append(json_response['results'][i]['pubDate'])
        img_url.append(json_response['results'][i]['image_url'])
        source.append(json_response['results'][i]['source_id'])

    data = zip(title, link, description, date, img_url, source)

    return render_template('sports.html', data=data)

# business page
@app.route('/business')
def business(): 

    title = []
    link = []
    description = []
    date = []
    img_url = []
    source = []


    response = requests.get("https://newsdata.io/api/1/news?apikey=pub_3357cb92f61b1a3819ff35e2ee398a56d2b3&language=en&category=business")
    json_response = response.json()

    for i in range(len(json_response['results'])):
        try:
            if json_response['results'][i]['full_description'] != None:
                description.append(json_response['results'][i]['full_description'])
            elif json_response['results'][i]['description'] != None:
                description.append(json_response['results'][i]['description'])
            else:
                continue
        except:
            continue
        title.append(json_response['results'][i]['title'])
        link.append(json_response['results'][i]['link'])
        date.append(json_response['results'][i]['pubDate'])
        img_url.append(json_response['results'][i]['image_url'])
        source.append(json_response['results'][i]['source_id'])

    data = zip(title, link, description, date, img_url, source)

    return render_template('business.html', data=data)

# save news in doc format
@app.route("/doc", methods=['GET'])
def doc():
    print("download as doc")
    document = Document()
    document.add_heading(request.args.get("title"))
    document.add_paragraph(request.args.get("news"))
    document.save('./static/files/saved_doc.docx')
    return send_from_directory(directory="./static/files", filename='saved_doc.docx')

# crypto page
@app.route('/crypto')
def crypto(): 

    title = []
    link = []
    description = []
    date = []
    img_url = []
    source = []


    response = requests.get("https://newsdata.io/api/1/news?apikey=pub_3357cb92f61b1a3819ff35e2ee398a56d2b3&language=en&q=cryptocurrency")
    json_response = response.json()

    for i in range(len(json_response['results'])):
        try:
            if json_response['results'][i]['full_description'] != None:
                description.append(json_response['results'][i]['full_description'])
            elif json_response['results'][i]['description'] != None:
                description.append(json_response['results'][i]['description'])
            else:
                continue
        except:
            continue
        title.append(json_response['results'][i]['title'])
        link.append(json_response['results'][i]['link'])
        date.append(json_response['results'][i]['pubDate'])
        img_url.append(json_response['results'][i]['image_url'])
        source.append(json_response['results'][i]['source_id'])

    data = zip(title, link, description, date, img_url, source)

    
    res = requests.get("https://www.tradingview.com/markets/cryptocurrencies/prices-all/")
    soup = BeautifulSoup(res.content, "html.parser")
     
    top3coins = []
    topcoins = soup.findAll("a", class_="tv-screener__symbol")
    for i in range(3):
        top3coins.append(topcoins[i].text)
    coinsprice = soup.findAll("td", class_="tv-data-table__cell tv-screener-table__cell tv-screener-table__cell--big")[:18]
    top3coinsprice = []
    top3coinsprice.append(coinsprice[2].text)
    top3coinsprice.append(coinsprice[8].text)
    top3coinsprice.append(coinsprice[14].text)

    return render_template('crypto.html', data=data, t3c=top3coins, t3cp=top3coinsprice)

# returns cypto data
def gettopcryptos():
    res = requests.get("https://www.tradingview.com/markets/cryptocurrencies/prices-all/")
    soup = BeautifulSoup(res.content, "html.parser")

    coin_name = []
    mkt_cap = []
    fd_mkt_cap = []
    last = []
    avail_coins = []
    total_coins = []
    traded_vol = []

    
    res = requests.get("https://www.tradingview.com/markets/cryptocurrencies/prices-all/")
    soup = BeautifulSoup(res.content, "html.parser")

    coin_name = []
    mkt_cap = []
    fd_mkt_cap = []
    last = []
    avail_coins = []
    total_coins = []
    traded_vol = []

    
    coins_len = len(soup.findAll("a", class_="tv-screener__symbol"))

    # coin name
    for i in range(coins_len):
        coin_name.append(soup.findAll("a", class_="tv-screener__symbol")[i].text)

    # other info (market cap, fully diluted market cap, last price, available coins, total coins, traded volume)
    other_info = []
    x, y = 0, 6
    for i in range(coins_len):
        temp = []
        for stat in soup.findAll("td", class_="tv-data-table__cell tv-screener-table__cell tv-screener-table__cell--big")[x:y]:
            temp.append(stat.text)
        x = y
        y += 6
        other_info.append(temp)

    for i in other_info:
        mkt_cap.append(i[0])
        fd_mkt_cap.append(i[1])
        last.append(i[2])
        avail_coins.append(i[3])
        total_coins.append(i[4])
        traded_vol.append(i[5])

    coin_info = []
    idx = 0
    for i in coin_name:
        info = []
        for j in other_info[idx]:
            info.append(j)
        coin_info.append([i] + info)
        idx += 1

    return coin_info

# top cryptos page
@app.route("/topcrypto")
def topcrypto():

    coin_info = gettopcryptos()

    return render_template('topcryptos.html', coin_data=coin_info)

# save top cryptos data in excel format
@app.route("/excel", methods=['GET'])
def excel():
    print("download as excel")

    coin_info = gettopcryptos()
    coins_len = len(coin_info)
        
    wb = Workbook()

    sheet = wb.active

    sheet['A1'] = 'Coin Name'
    sheet['B1'] = 'Market Cap'
    sheet['C1'] = 'Fully Diluted Market Cap'
    sheet['D1'] = 'Last Price'
    sheet['E1'] = 'Available Coins'
    sheet['F1'] = 'Total Coins'
    sheet['G1'] = 'Traded Volume'

    row = 2

    for i in range(coins_len):
        sheet['A'+str(row)] = coin_info[i][0]
        sheet['B'+str(row)] = coin_info[i][1]
        sheet['C'+str(row)] = coin_info[i][2]
        sheet['D'+str(row)] = coin_info[i][3]
        sheet['E'+str(row)] = coin_info[i][4]
        sheet['F'+str(row)] = coin_info[i][5]
        sheet['G'+str(row)] = coin_info[i][6]
        row += 1

    wb.save("./static/files/Crypto_Data.xlsx")

    return send_from_directory(directory="./static/files", filename='Crypto_Data.xlsx')

# save top cryptos data in csv
@app.route("/csv", methods=['GET'])
def csv_():
    print("download as csv")

    coin_info = gettopcryptos()
    coins_len = len(coin_info)

    with open("./static/files/Crypto_Data.csv" , "w") as file:
        writer = csv.writer(file)
        writer.writerow(['Coin name', 'Market Cap', 'Fully Diluted Market Cap', 'Last Price', 'Available Coins', 'Total Coins', 'Traded Volume'])
        writer.writerows(coin_info)

    return send_from_directory(directory="./static/files", filename='Crypto_Data.csv')

# API

# put request parser
crypto_put_args = reqparse.RequestParser()
crypto_put_args.add_argument('crypto-name', type=str, required=True)
crypto_put_args.add_argument('currency-name', type=str, required=True)

class api_topcryptos(Resource):

    def get(self):
        coin_info = gettopcryptos()
        json_dict = {}
        for info in coin_info:
            stats = {}
            stats['mkt_cap'] = info[1]
            stats['fd_mkt_cap'] = info[2]
            stats['last'] = info[3]
            stats['avail_coins'] = info[4]
            stats['total_coins'] = info[5]
            stats['traded_vol'] = info[6]
            json_dict[info[0]] = stats
        json_response = json.dumps(json_dict)
        return json_response

    def put(self):
        args = crypto_put_args.parse_args()
        crypto_name = args['crypto-name']
        currency_name = args['currency-name']
        print(crypto_name)
        print(currency_name)
        res = requests.get(f"https://api.coingecko.com/api/v3/simple/price?ids={crypto_name}&vs_currencies={currency_name}")
        return res.json(), 200

api.add_resource(api_topcryptos, '/api/cryptodata')
    
        
if __name__ == '__main__':
    app.run(debug=True)
   

